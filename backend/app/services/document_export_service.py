from __future__ import annotations

from dataclasses import dataclass
from io import BytesIO
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt
from reportlab.lib.enums import TA_JUSTIFY, TA_LEFT
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer
from sqlalchemy.orm import Session

from backend.app.models import Case, DebtorProfile
from backend.app.services.case_service import get_projection_data_or_rebuild
from backend.app.services.legal_document_templates import build_legal_document_template


@dataclass
class ExportedDocument:
    filename: str
    media_type: str
    content: bytes


def _load_case_or_raise(db: Session, case_id: int) -> Case:
    case = db.query(Case).filter(Case.id == case_id).first()
    if not case:
        raise ValueError(f"Case {case_id} not found")
    return case


def _load_debtor_profile(db: Session, case_id: int) -> DebtorProfile | None:
    return db.query(DebtorProfile).filter(DebtorProfile.case_id == case_id).first()


def _build_template(db: Session, case_id: int, document_code: str) -> dict:
    case = _load_case_or_raise(db, case_id)
    projection = get_projection_data_or_rebuild(db, case_id)
    debtor_profile = _load_debtor_profile(db, case_id)

    return build_legal_document_template(
        db=db,
        tenant_id=case.tenant_id,
        document_code=document_code,
        case=case,
        projection=projection,
        debtor_profile=debtor_profile,
    )


def _apply_docx_default_font(document: Document) -> None:
    styles = document.styles

    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")

    if "Title" in styles:
        title = styles["Title"]
        title.font.name = "Times New Roman"
        title.font.size = Pt(14)
        title._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")


def export_document_docx(
    db: Session,
    *,
    case_id: int,
    document_code: str,
) -> ExportedDocument:
    template = _build_template(db, case_id, document_code)

    document = Document()
    _apply_docx_default_font(document)

    title = document.add_paragraph()
    title.style = document.styles["Title"]
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title_run = title.add_run(template["title"])
    title_run.bold = True
    title_run.font.name = "Times New Roman"
    title_run.font.size = Pt(14)
    title_run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")

    document.add_paragraph("")

    for paragraph in template["paragraphs"]:
        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY if paragraph else WD_ALIGN_PARAGRAPH.LEFT

        run = p.add_run(paragraph if paragraph else "")
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")

    output = BytesIO()
    document.save(output)

    filename = f"case_{case_id}_{template['file_stub']}.docx"
    return ExportedDocument(
        filename=filename,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        content=output.getvalue(),
    )


def _pick_pdf_font() -> str:
    candidates = [
        ("DejaVuSerif", "/usr/share/fonts/truetype/dejavu/DejaVuSerif.ttf"),
        ("LiberationSerif", "/usr/share/fonts/truetype/liberation2/LiberationSerif-Regular.ttf"),
    ]

    registered = set(pdfmetrics.getRegisteredFontNames())

    for font_name, font_path in candidates:
        if font_name in registered:
            return font_name

        if Path(font_path).exists():
            try:
                pdfmetrics.registerFont(TTFont(font_name, font_path))
                return font_name
            except Exception:
                continue

    return "Times-Roman"


def export_document_pdf(
    db: Session,
    *,
    case_id: int,
    document_code: str,
) -> ExportedDocument:
    template = _build_template(db, case_id, document_code)
    font_name = _pick_pdf_font()

    output = BytesIO()
    doc = SimpleDocTemplate(
        output,
        pagesize=A4,
        leftMargin=20 * mm,
        rightMargin=20 * mm,
        topMargin=18 * mm,
        bottomMargin=18 * mm,
    )

    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        "DebtrixTitle",
        parent=styles["Heading1"],
        fontName=font_name,
        fontSize=14,
        leading=18,
        alignment=TA_LEFT,
        spaceAfter=10,
    )
    body_style = ParagraphStyle(
        "DebtrixBody",
        parent=styles["BodyText"],
        fontName=font_name,
        fontSize=11,
        leading=15,
        alignment=TA_JUSTIFY,
        spaceAfter=7,
    )

    story = [Paragraph(template["title"], title_style), Spacer(1, 4)]

    for paragraph in template["paragraphs"]:
        if not paragraph:
            story.append(Spacer(1, 8))
            continue

        safe_text = (
            paragraph.replace("&", "&amp;")
            .replace("<", "&lt;")
            .replace(">", "&gt;")
        )
        story.append(Paragraph(safe_text, body_style))

    doc.build(story)

    filename = f"case_{case_id}_{template['file_stub']}.pdf"
    return ExportedDocument(
        filename=filename,
        media_type="application/pdf",
        content=output.getvalue(),
    )